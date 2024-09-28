import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter import PhotoImage
import requests
import os
import PyPDF2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START
import re
import json
from datetime import datetime

class ScriptsWindow(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.title("Manage Scripts")
        self.geometry("500x400")
        self.attributes('-topmost', True)
        self.create_widgets()

    def create_widgets(self):
        self.scripts_listbox = tk.Listbox(self, width=70, height=15)
        self.scripts_listbox.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

        buttons_frame = ttk.Frame(self)
        buttons_frame.pack(pady=10, fill=tk.X)

        upload_btn = ttk.Button(buttons_frame, text="Upload Script", command=self.upload_script)
        upload_btn.pack(side=tk.LEFT, padx=5)

        move_up_btn = ttk.Button(buttons_frame, text="Move Up", command=self.move_up)
        move_up_btn.pack(side=tk.LEFT, padx=5)

        move_down_btn = ttk.Button(buttons_frame, text="Move Down", command=self.move_down)
        move_down_btn.pack(side=tk.LEFT, padx=5)

        delete_btn = ttk.Button(buttons_frame, text="Delete Selected", command=self.delete_selected)
        delete_btn.pack(side=tk.LEFT, padx=5)

        save_btn = ttk.Button(buttons_frame, text="Save", command=self.parent.save_settings)
        save_btn.pack(side=tk.LEFT, padx=5)

        close_btn = ttk.Button(buttons_frame, text="Close", command=self.destroy)
        close_btn.pack(side=tk.RIGHT, padx=5)

        self.update_listbox()

    def upload_script(self):
        file_types = [("PDF files", "*.pdf"), ("Text files", "*.txt"), ("All files", "*.*")]
        file_paths = filedialog.askopenfilenames(title="Select Script(s) or Paper(s)", filetypes=file_types)
        for file_path in file_paths:
            self.parent.upload_script(file_path)
        self.update_listbox()

    def move_up(self):
        selection = self.scripts_listbox.curselection()
        if selection:
            index = selection[0]
            if index > 0:
                self.parent.scripts[index], self.parent.scripts[index - 1] = self.parent.scripts[index - 1], self.parent.scripts[index]
                self.update_listbox()
                self.scripts_listbox.select_set(index - 1)

    def move_down(self):
        selection = self.scripts_listbox.curselection()
        if selection:
            index = selection[0]
            if index < len(self.parent.scripts) - 1:
                self.parent.scripts[index], self.parent.scripts[index + 1] = self.parent.scripts[index + 1], self.parent.scripts[index]
                self.update_listbox()
                self.scripts_listbox.select_set(index + 1)

    def delete_selected(self):
        selection = self.scripts_listbox.curselection()
        if selection:
            index = selection[0]
            del self.parent.scripts[index]
            self.update_listbox()

    def update_listbox(self):
        self.scripts_listbox.delete(0, tk.END)
        for script in self.parent.scripts:
            self.scripts_listbox.insert(tk.END, script[0])

class InstructionsWindow(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.title("Manage Instructions")
        self.geometry("500x400")
        self.attributes('-topmost', True)
        self.create_widgets()

    def create_widgets(self):
        self.instructions_listbox = tk.Listbox(self, width=70, height=15)
        self.instructions_listbox.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

        buttons_frame = ttk.Frame(self)
        buttons_frame.pack(pady=10, fill=tk.X)

        upload_btn = ttk.Button(buttons_frame, text="Upload Instruction", command=self.upload_instruction)
        upload_btn.pack(side=tk.LEFT, padx=5)

        move_up_btn = ttk.Button(buttons_frame, text="Move Up", command=self.move_up)
        move_up_btn.pack(side=tk.LEFT, padx=5)

        move_down_btn = ttk.Button(buttons_frame, text="Move Down", command=self.move_down)
        move_down_btn.pack(side=tk.LEFT, padx=5)

        delete_btn = ttk.Button(buttons_frame, text="Delete Selected", command=self.delete_selected)
        delete_btn.pack(side=tk.LEFT, padx=5)

        save_btn = ttk.Button(buttons_frame, text="Save", command=self.parent.save_settings)
        save_btn.pack(side=tk.LEFT, padx=5)

        close_btn = ttk.Button(buttons_frame, text="Close", command=self.destroy)
        close_btn.pack(side=tk.RIGHT, padx=5)

        self.update_listbox()

    def upload_instruction(self):
        file_types = [("PDF files", "*.pdf"), ("Text files", "*.txt"), ("All files", "*.*")]
        file_paths = filedialog.askopenfilenames(title="Select Instruction File(s)", filetypes=file_types)
        for file_path in file_paths:
            self.parent.upload_instruction(file_path)
        self.update_listbox()

    def move_up(self):
        selection = self.instructions_listbox.curselection()
        if selection:
            index = selection[0]
            if index > 0:
                self.parent.instructions[index], self.parent.instructions[index - 1] = self.parent.instructions[index - 1], self.parent.instructions[index]
                self.update_listbox()
                self.instructions_listbox.select_set(index - 1)

    def move_down(self):
        selection = self.instructions_listbox.curselection()
        if selection:
            index = selection[0]
            if index < len(self.parent.instructions) - 1:
                self.parent.instructions[index], self.parent.instructions[index + 1] = self.parent.instructions[index + 1], self.parent.instructions[index]
                self.update_listbox()
                self.instructions_listbox.select_set(index + 1)

    def delete_selected(self):
        selection = self.instructions_listbox.curselection()
        if selection:
            index = selection[0]
            del self.parent.instructions[index]
            self.update_listbox()

    def update_listbox(self):
        self.instructions_listbox.delete(0, tk.END)
        for instruction in self.parent.instructions:
            self.instructions_listbox.insert(tk.END, instruction[0])

class SettingsWindow(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.title("Settings")
        self.geometry("400x300")
        self.attributes('-topmost', True)
        self.create_widgets()

    def create_widgets(self):
        main_frame = ttk.Frame(self, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # API Key
        ttk.Label(main_frame, text="API Key:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.api_key_entry = ttk.Entry(main_frame, width=40, show='*')
        self.api_key_entry.grid(row=0, column=1, pady=5, padx=5, sticky=tk.W+tk.E)
        self.api_key_entry.insert(0, self.parent.api_key)

        # First Name
        ttk.Label(main_frame, text="First Name:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.first_name_entry = ttk.Entry(main_frame, width=40)
        self.first_name_entry.grid(row=1, column=1, pady=5, padx=5, sticky=tk.W+tk.E)
        self.first_name_entry.insert(0, self.parent.first_name)

        # Last Name
        ttk.Label(main_frame, text="Last Name:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.last_name_entry = ttk.Entry(main_frame, width=40)
        self.last_name_entry.grid(row=2, column=1, pady=5, padx=5, sticky=tk.W+tk.E)
        self.last_name_entry.insert(0, self.parent.last_name)

        # Date
        ttk.Label(main_frame, text="Date:").grid(row=3, column=0, sticky=tk.W, pady=5)
        self.date_entry = ttk.Entry(main_frame, width=40)
        self.date_entry.grid(row=3, column=1, pady=5, padx=5, sticky=tk.W+tk.E)
        self.date_entry.insert(0, self.parent.date)

        # Buttons
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.grid(row=4, column=0, columnspan=2, pady=20)

        save_btn = ttk.Button(buttons_frame, text="Save", command=self.save_settings)
        save_btn.pack(side=tk.LEFT, padx=5)

        close_btn = ttk.Button(buttons_frame, text="Close", command=self.destroy)
        close_btn.pack(side=tk.RIGHT, padx=5)

    def save_settings(self):
        self.parent.api_key = self.api_key_entry.get()
        self.parent.first_name = self.first_name_entry.get()
        self.parent.last_name = self.last_name_entry.get()
        self.parent.date = self.date_entry.get()
        self.parent.save_settings()
        messagebox.showinfo("Settings", "Settings saved successfully.")

class FormattingWindow(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.title("Formatting Options")
        self.geometry("400x500")
        self.attributes('-topmost', True)
        self.create_widgets()

    def create_widgets(self):
        main_frame = ttk.Frame(self, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Font Name
        ttk.Label(main_frame, text="Font Name:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.font_name_var = tk.StringVar(value=self.parent.font_name)
        font_options = ["Times New Roman", "Arial", "Calibri", "Cambria", "Verdana"]
        self.font_name_combobox = ttk.Combobox(main_frame, textvariable=self.font_name_var, values=font_options, state="readonly")
        self.font_name_combobox.grid(row=0, column=1, sticky=tk.W+tk.E, pady=5, padx=5)

        # Line Spacing
        ttk.Label(main_frame, text="Line Spacing:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.line_spacing_var = tk.StringVar(value=self.parent.line_spacing)
        line_spacing_options = ["Single", "1.5 lines", "Double"]
        self.line_spacing_combobox = ttk.Combobox(main_frame, textvariable=self.line_spacing_var, values=line_spacing_options, state="readonly")
        self.line_spacing_combobox.grid(row=1, column=1, sticky=tk.W+tk.E, pady=5, padx=5)

        # Font Sizes
        ttk.Label(main_frame, text="Font Sizes (pt):").grid(row=2, column=0, sticky=tk.W, pady=5)
        sizes_frame = ttk.Frame(main_frame)
        sizes_frame.grid(row=2, column=1, sticky=tk.W+tk.E, pady=5, padx=5)
        sizes_frame.columnconfigure(1, weight=1)

        ttk.Label(sizes_frame, text="Normal Text:").grid(row=0, column=0, sticky=tk.W)
        self.font_size_normal_var = tk.IntVar(value=self.parent.font_size_normal)
        self.font_size_normal_spinbox = ttk.Spinbox(sizes_frame, from_=8, to=72, textvariable=self.font_size_normal_var)
        self.font_size_normal_spinbox.grid(row=0, column=1, sticky=tk.W+tk.E)

        ttk.Label(sizes_frame, text="Heading 1:").grid(row=1, column=0, sticky=tk.W)
        self.font_size_heading1_var = tk.IntVar(value=self.parent.font_size_heading1)
        self.font_size_heading1_spinbox = ttk.Spinbox(sizes_frame, from_=8, to=72, textvariable=self.font_size_heading1_var)
        self.font_size_heading1_spinbox.grid(row=1, column=1, sticky=tk.W+tk.E)

        ttk.Label(sizes_frame, text="Heading 2:").grid(row=2, column=0, sticky=tk.W)
        self.font_size_heading2_var = tk.IntVar(value=self.parent.font_size_heading2)
        self.font_size_heading2_spinbox = ttk.Spinbox(sizes_frame, from_=8, to=72, textvariable=self.font_size_heading2_var)
        self.font_size_heading2_spinbox.grid(row=2, column=1, sticky=tk.W+tk.E)

        ttk.Label(sizes_frame, text="Heading 3:").grid(row=3, column=0, sticky=tk.W)
        self.font_size_heading3_var = tk.IntVar(value=self.parent.font_size_heading3)
        self.font_size_heading3_spinbox = ttk.Spinbox(sizes_frame, from_=8, to=72, textvariable=self.font_size_heading3_var)
        self.font_size_heading3_spinbox.grid(row=3, column=1, sticky=tk.W+tk.E)

        # Margins
        ttk.Label(main_frame, text="Margins (cm):").grid(row=3, column=0, sticky=tk.W, pady=5)
        margins_frame = ttk.Frame(main_frame)
        margins_frame.grid(row=3, column=1, sticky=tk.W+tk.E, pady=5, padx=5)
        margins_frame.columnconfigure(1, weight=1)

        ttk.Label(margins_frame, text="Top:").grid(row=0, column=0, sticky=tk.W)
        self.margin_top_var = tk.DoubleVar(value=self.parent.margins.get('top', 2.0))
        self.margin_top_spinbox = ttk.Spinbox(margins_frame, from_=0.5, to=5.0, increment=0.1, textvariable=self.margin_top_var)
        self.margin_top_spinbox.grid(row=0, column=1, sticky=tk.W+tk.E, pady=2)

        ttk.Label(margins_frame, text="Bottom:").grid(row=1, column=0, sticky=tk.W)
        self.margin_bottom_var = tk.DoubleVar(value=self.parent.margins.get('bottom', 2.0))
        self.margin_bottom_spinbox = ttk.Spinbox(margins_frame, from_=0.5, to=5.0, increment=0.1, textvariable=self.margin_bottom_var)
        self.margin_bottom_spinbox.grid(row=1, column=1, sticky=tk.W+tk.E, pady=2)

        ttk.Label(margins_frame, text="Left:").grid(row=2, column=0, sticky=tk.W)
        self.margin_left_var = tk.DoubleVar(value=self.parent.margins.get('left', 2.0))
        self.margin_left_spinbox = ttk.Spinbox(margins_frame, from_=0.5, to=5.0, increment=0.1, textvariable=self.margin_left_var)
        self.margin_left_spinbox.grid(row=2, column=1, sticky=tk.W+tk.E, pady=2)

        ttk.Label(margins_frame, text="Right:").grid(row=3, column=0, sticky=tk.W)
        self.margin_right_var = tk.DoubleVar(value=self.parent.margins.get('right', 2.0))
        self.margin_right_spinbox = ttk.Spinbox(margins_frame, from_=0.5, to=5.0, increment=0.1, textvariable=self.margin_right_var)
        self.margin_right_spinbox.grid(row=3, column=1, sticky=tk.W+tk.E, pady=2)

        # Buttons
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.grid(row=4, column=0, columnspan=2, pady=20)

        save_btn = ttk.Button(buttons_frame, text="Save", command=self.save_settings)
        save_btn.pack(side=tk.LEFT, padx=5)

        default_btn = ttk.Button(buttons_frame, text="Restore Defaults", command=self.restore_defaults)
        default_btn.pack(side=tk.LEFT, padx=5)

        close_btn = ttk.Button(buttons_frame, text="Close", command=self.destroy)
        close_btn.pack(side=tk.RIGHT, padx=5)

    def save_settings(self):
        self.parent.font_name = self.font_name_var.get()
        self.parent.line_spacing = self.line_spacing_var.get()
        self.parent.font_size_normal = self.font_size_normal_var.get()
        self.parent.font_size_heading1 = self.font_size_heading1_var.get()
        self.parent.font_size_heading2 = self.font_size_heading2_var.get()
        self.parent.font_size_heading3 = self.font_size_heading3_var.get()
        self.parent.margins = {
            'top': self.margin_top_var.get(),
            'bottom': self.margin_bottom_var.get(),
            'left': self.margin_left_var.get(),
            'right': self.margin_right_var.get()
        }
        self.parent.save_settings()
        messagebox.showinfo("Formatting", "Formatting settings saved successfully.")

    def restore_defaults(self):
        self.font_name_var.set("Times New Roman")
        self.line_spacing_var.set("1.5 lines")
        self.font_size_normal_var.set(12)
        self.font_size_heading1_var.set(16)
        self.font_size_heading2_var.set(14)
        self.font_size_heading3_var.set(12)
        self.margin_top_var.set(2.0)
        self.margin_bottom_var.set(2.0)
        self.margin_left_var.set(2.0)
        self.margin_right_var.set(2.0)
        messagebox.showinfo("Formatting", "Formatting settings restored to defaults.")

class ClaudeApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("University Paper Generator")
        self.geometry("800x800")

        # Initialize variables
        self.api_key = ""
        self.system_prompt = (
            "You are to write a university paper based on the provided scientific papers and study scripts. "
            "Determine an appropriate title for the paper. The paper should be formatted as a real university paper suitable for submission, "
            "including chapters, sections, headings, and citations in Harvard style. "
            "Include a bibliography at the end. Do not include a table of contents. "
            "Take content from the scripts provided below and cite them appropriately. "
            "Write in first person singular, as if by {first_name} {last_name}. "
            "The paper should be dated {date}. "
            "Remember to sometimes make minor spelling mistakes as a student around B2/C1 English level might. "
            "Your specific instructions are:\n\n"
            "{instructions}\n\n"
            "!!!this is the next document!!!\n\n"
            "{scripts}\n\n"
            "Please output the paper in Markdown format with clear markers for headings and sections. "
            "Use '#' for main headings, '##' for subheadings, and '###' for sub-subheadings. "
            "Use **bold** and *italic* text where appropriate. Include bullet points and numbered lists if necessary. "
            "Ensure that citations are properly formatted in Harvard style and included within the text. "
            "At the beginning of the paper, include a title page containing the paper's title, your name, and date. "
            "Enclose the title page content between '####TITLE PAGE####' and '####END TITLE PAGE####'.\n\n"
            "The structure of the paper should be:\n\n"
            "####TITLE PAGE####\n"
            "# [Title of the Paper]\n"
            "Author: {first_name} {last_name}\n"
            "Date: {date}\n"
            "####END TITLE PAGE####\n\n"
            "# Introduction\n"
            "...\n\n"
            "# Conclusion\n"
            "...\n\n"
            "# Bibliography\n"
            "...\n\n"
            "IMPORTANT: IF THE INSTRUCTIONS OR DETAILS SUCH AS TITLE PAGE, "
            "BIBLIOGRAPHY, CITATION STYLE, ETC., ARE PROVIDED REGARDING STRUCTURING THE PAPER, "
            "PLEASE FOLLOW THOSE INSTEAD OF THE ONES LISTED ABOVE. MAKE USE OF FULL MAX TOKEN OUTPUT OF 8192"
        )
        self.scripts = []
        self.instructions = []
        self.first_name = ""
        self.last_name = ""
        self.date = datetime.now().strftime("%Y-%m-%d")
        self.margins = {'top': 2.0, 'bottom': 2.0, 'left': 2.0, 'right': 2.0}

        # Formatting options
        self.font_name = "Times New Roman"
        self.font_size_normal = 12
        self.font_size_heading1 = 16
        self.font_size_heading2 = 14
        self.font_size_heading3 = 12
        self.line_spacing = 1.5  # Default line spacing

        # Load saved settings
        self.load_settings()

        # Build UI
        self.create_widgets()

    def create_widgets(self):
        style = ttk.Style()
        style.theme_use('clam')

        # Top Frame for Settings and Formatting Buttons
        top_frame = ttk.Frame(self, padding="10")
        top_frame.pack(fill=tk.X)

        settings_icon = PhotoImage(width=20, height=20)  # Placeholder for gear icon
        try:
            settings_icon = PhotoImage(file="gear.png")  # Ensure you have a gear.png in the directory
        except:
            pass  # Use placeholder if gear.png not found

        settings_btn = ttk.Button(top_frame, text="Settings", image=settings_icon, compound=tk.LEFT, command=self.open_settings_window)
        settings_btn.image = settings_icon  # Keep a reference
        settings_btn.pack(side=tk.RIGHT, padx=5)

        formatting_btn = ttk.Button(top_frame, text="Formatting", command=self.open_formatting_window)
        formatting_btn.pack(side=tk.RIGHT, padx=5)

        # Main Frame
        main_frame = ttk.Frame(self, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # API Key (Moved to Settings)
        # First Name (Moved to Settings)
        # Last Name (Moved to Settings)
        # Date (Moved to Settings)

        # Instructions Button
        instructions_btn = ttk.Button(main_frame, text="Manage Instructions", command=self.open_instructions_window)
        instructions_btn.grid(row=0, column=0, columnspan=3, pady=10, sticky=tk.W+tk.E)

        # Scripts Button
        scripts_btn = ttk.Button(main_frame, text="Manage Scripts", command=self.open_scripts_window)
        scripts_btn.grid(row=1, column=0, columnspan=3, pady=10, sticky=tk.W+tk.E)

        # System Prompt (Read-Only)
        ttk.Label(main_frame, text="System Prompt:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.system_prompt_text = tk.Text(main_frame, wrap=tk.WORD, height=10, state='disabled')
        self.system_prompt_text.grid(row=3, column=0, columnspan=3, pady=5, padx=5, sticky=tk.W+tk.E+tk.N+tk.S)
        self.update_system_prompt_text()

        # Buttons
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.grid(row=4, column=0, columnspan=3, pady=10)

        generate_btn = ttk.Button(buttons_frame, text="Generate Paper", command=self.send_request)
        generate_btn.pack(side=tk.LEFT, padx=5)

        save_output_btn = ttk.Button(buttons_frame, text="Save Output", command=self.save_output)
        save_output_btn.pack(side=tk.LEFT, padx=5)

        # Output Text Area
        self.output_text = tk.Text(main_frame, wrap=tk.WORD, height=15)
        self.output_text.grid(row=5, column=0, columnspan=3, pady=10, padx=5, sticky=tk.W+tk.E+tk.N+tk.S)

        # Configure grid weights
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(3, weight=1)
        main_frame.rowconfigure(5, weight=2)

    def load_settings(self):
        try:
            with open('claude_app_settings.json', 'r') as f:
                settings = json.load(f)
            self.api_key = settings.get('api_key', '')
            self.first_name = settings.get('first_name', '')
            self.last_name = settings.get('last_name', '')
            self.system_prompt = settings.get('system_prompt', self.system_prompt)
            self.font_name = settings.get('font_name', 'Times New Roman')
            self.font_size_normal = settings.get('font_size_normal', 12)
            self.font_size_heading1 = settings.get('font_size_heading1', 16)
            self.font_size_heading2 = settings.get('font_size_heading2', 14)
            self.font_size_heading3 = settings.get('font_size_heading3', 12)
            self.line_spacing = settings.get('line_spacing', '1.5 lines')
            self.margins = settings.get('margins', {'top': 2.0, 'bottom': 2.0, 'left': 2.0, 'right': 2.0})
        except FileNotFoundError:
            pass  # It's okay if the file doesn't exist yet

    def save_settings(self):
        settings = {
            'api_key': self.api_key,
            'first_name': self.first_name,
            'last_name': self.last_name,
            'system_prompt': self.system_prompt,
            'font_name': self.font_name,
            'font_size_normal': self.font_size_normal,
            'font_size_heading1': self.font_size_heading1,
            'font_size_heading2': self.font_size_heading2,
            'font_size_heading3': self.font_size_heading3,
            'line_spacing': self.line_spacing,
            'margins': self.margins
        }
        with open('claude_app_settings.json', 'w') as f:
            json.dump(settings, f)
        self.update_system_prompt_text()

    def update_system_prompt_text(self):
        self.system_prompt_text.config(state='normal')
        formatted_instructions = "\n\n".join([f"Instruction {i+1} ({name}):\n{content}" for i, (name, content) in enumerate(self.instructions)])
        formatted_scripts = "\n\n".join([f"Script {i+1} ({name}):\n{content}" for i, (name, content) in enumerate(self.scripts)])
        system_message = self.system_prompt.format(
            scripts=formatted_scripts,
            instructions=formatted_instructions,
            first_name=self.first_name,
            last_name=self.last_name,
            date=self.date
        )
        self.system_prompt_text.delete(1.0, tk.END)
        self.system_prompt_text.insert(tk.END, system_message)
        self.system_prompt_text.config(state='disabled')

    def open_scripts_window(self):
        ScriptsWindow(self)

    def open_instructions_window(self):
        InstructionsWindow(self)

    def open_settings_window(self):
        SettingsWindow(self)

    def open_formatting_window(self):
        FormattingWindow(self)

    def upload_script(self, file_path):
        file_name = os.path.basename(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.pdf':
            # Extract text from PDF
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
                self.scripts.append((file_name, text))
            except Exception as e:
                messagebox.showerror("Error", f"Error reading PDF file {file_name}: {e}")
        else:
            # Read text file
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                self.scripts.append((file_name, text))
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        text = f.read()
                    self.scripts.append((file_name, text))
                except Exception as e:
                    messagebox.showerror("Error", f"Error reading text file {file_name}: {e}")
        self.update_system_prompt_text()

    def upload_instruction(self, file_path):
        file_name = os.path.basename(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.pdf':
            # Extract text from PDF
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
                self.instructions.append((file_name, text))
            except Exception as e:
                messagebox.showerror("Error", f"Error reading PDF file {file_name}: {e}")
        else:
            # Read text file
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                self.instructions.append((file_name, text))
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        text = f.read()
                    self.instructions.append((file_name, text))
                except Exception as e:
                    messagebox.showerror("Error", f"Error reading text file {file_name}: {e}")
        self.update_system_prompt_text()

    def send_request(self):
        if not self.instructions:
            messagebox.showerror("Error", "Please upload instruction files first.")
            return
        if not self.scripts:
            messagebox.showerror("Error", "Please upload script files first.")
            return
        if not self.first_name or not self.last_name:
            messagebox.showerror("Error", "Please enter your first and last name in Settings.")
            return
        if not self.api_key:
            messagebox.showerror("Error", "Please enter your API key in Settings.")
            return

        self.first_name = self.first_name
        self.last_name = self.last_name
        self.date = self.date
        self.save_settings()

        formatted_instructions = "\n\n".join([f"Instruction {i+1} ({name}):\n{content}" for i, (name, content) in enumerate(self.instructions)])
        formatted_scripts = "\n\n".join([f"Script {i+1} ({name}):\n{content}" for i, (name, content) in enumerate(self.scripts)])

        system_message = self.system_prompt.format(
            scripts=formatted_scripts,
            instructions=formatted_instructions,
            first_name=self.first_name,
            last_name=self.last_name,
            date=self.date
        )

        messages = [
            {"role": "user", "content": system_message}
        ]

        # API call parameters
        api_url = "https://api.anthropic.com/v1/messages"
        headers = {
            "x-api-key": self.api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
            "anthropic-beta": "max-tokens-3-5-sonnet-2024-07-15"
        }
        data = {
            "model": "claude-3-5-sonnet-20240620",
            "max_tokens": 8192,
            "messages": messages
        }

        try:
            # Show a loading message
            self.output_text.delete(1.0, tk.END)
            self.output_text.insert(tk.END, "Generating response, please wait...")
            self.update_idletasks()

            response = requests.post(api_url, headers=headers, json=data)
            if response.status_code == 200:
                result = response.json()
                content = result['content'][0]['text']
                response_text = content.strip()
                self.output_text.delete(1.0, tk.END)
                self.output_text.insert(tk.END, response_text)
                messagebox.showinfo("Success", "Paper generated.")
            else:
                error_message = response.text
                self.output_text.delete(1.0, tk.END)
                self.output_text.insert(tk.END, "")
                messagebox.showerror("Error", f"API Error {response.status_code}: {error_message}")
        except Exception as e:
            self.output_text.delete(1.0, tk.END)
            messagebox.showerror("Error", f"Error making API request: {e}")

    def save_output(self):
        output = self.output_text.get(1.0, tk.END).strip()
        if not output:
            messagebox.showerror("Error", "No output to save.")
            return

        save_path = filedialog.asksaveasfilename(
            title="Save Output as Word File",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        if save_path:
            try:
                # Get formatting options
                font_name = self.font_name
                font_size_normal = self.font_size_normal
                font_size_heading1 = self.font_size_heading1
                font_size_heading2 = self.font_size_heading2
                font_size_heading3 = self.font_size_heading3
                line_spacing_option = self.line_spacing
                if line_spacing_option == "Single":
                    line_spacing = 1
                elif line_spacing_option == "1.5 lines":
                    line_spacing = 1.5
                elif line_spacing_option == "Double":
                    line_spacing = 2
                else:
                    line_spacing = 1  # Default to single spacing

                # Document creation
                document = Document()

                # Document settings
                section = document.sections[0]
                section.page_height = Inches(11)
                section.page_width = Inches(8.5)
                section.left_margin = Inches(self.margins['left'] / 2.54)
                section.right_margin = Inches(self.margins['right'] / 2.54)
                section.top_margin = Inches(self.margins['top'] / 2.54)
                section.bottom_margin = Inches(self.margins['bottom'] / 2.54)

                # Define styles
                styles = document.styles

                # Title style
                if 'TitleStyle' not in styles:
                    style_title = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
                else:
                    style_title = styles['TitleStyle']
                style_title.font.size = Pt(font_size_heading1)
                style_title.font.bold = True
                style_title.font.name = font_name

                # Heading 1 style
                if 'Heading1Custom' not in styles:
                    style_heading1 = styles.add_style('Heading1Custom', WD_STYLE_TYPE.PARAGRAPH)
                else:
                    style_heading1 = styles['Heading1Custom']
                style_heading1.base_style = styles['Heading 1']
                style_heading1.font.size = Pt(font_size_heading1)
                style_heading1.font.bold = True
                style_heading1.font.name = font_name

                # Heading 2 style
                if 'Heading2Custom' not in styles:
                    style_heading2 = styles.add_style('Heading2Custom', WD_STYLE_TYPE.PARAGRAPH)
                else:
                    style_heading2 = styles['Heading2Custom']
                style_heading2.base_style = styles['Heading 2']
                style_heading2.font.size = Pt(font_size_heading2)
                style_heading2.font.bold = True
                style_heading2.font.name = font_name

                # Heading 3 style
                if 'Heading3Custom' not in styles:
                    style_heading3 = styles.add_style('Heading3Custom', WD_STYLE_TYPE.PARAGRAPH)
                else:
                    style_heading3 = styles['Heading3Custom']
                style_heading3.base_style = styles['Heading 3']
                style_heading3.font.size = Pt(font_size_heading3)
                style_heading3.font.bold = True
                style_heading3.font.name = font_name

                # Normal text style
                style_normal = styles['Normal']
                style_normal.font.size = Pt(font_size_normal)
                style_normal.font.name = font_name

                # Set line spacing for paragraph styles
                for style in [style_normal, style_title, style_heading1, style_heading2, style_heading3]:
                    style.paragraph_format.line_spacing = line_spacing

                # Extract title page content
                title_page_match = re.search(r'####TITLE PAGE####(.*?)####END TITLE PAGE####', output, re.DOTALL | re.IGNORECASE)
                if title_page_match:
                    title_page_content = title_page_match.group(1).strip()
                    output = output.replace(title_page_match.group(0), '')  # Remove title page from output

                    # Create title page
                    p = document.add_paragraph('', style='TitleStyle')
                    for line in title_page_content.split('\n'):
                        if line.strip():
                            p.add_run(line.strip()).add_break()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Start a new section on the next page
                    new_section = document.add_section(WD_SECTION_START.NEW_PAGE)
                else:
                    messagebox.showwarning("Warning", "Title page markers not found. Title page will not be created.")

                # Process the rest of the content
                paragraphs = output.strip().split('\n')
                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue
                    if para.startswith('# '):
                        # Heading level 1
                        document.add_paragraph(para[2:].strip(), style='Heading1Custom')
                    elif para.startswith('## '):
                        # Heading level 2
                        document.add_paragraph(para[3:].strip(), style='Heading2Custom')
                    elif para.startswith('### '):
                        # Heading level 3
                        document.add_paragraph(para[4:].strip(), style='Heading3Custom')
                    elif re.match(r'^\d+\.', para):
                        # Numbered list
                        document.add_paragraph(para, style='List Number')
                    elif para.startswith('- '):
                        # Bullet list
                        document.add_paragraph(para[2:].strip(), style='List Bullet')
                    else:
                        # Regular paragraph with bold and italic
                        p = document.add_paragraph(style='Normal')
                        self._add_runs(p, para)

                # Add page numbers
                for section in document.sections:
                    footer = section.footer
                    paragraph = footer.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    fldSimple = OxmlElement('w:fldSimple')
                    fldSimple.set(qn('w:instr'), 'PAGE')
                    run._r.append(fldSimple)

                # Save the document
                document.save(save_path)
                messagebox.showinfo("Success", f"Output saved to {save_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Error saving Word file: {e}")

    def _add_runs(self, paragraph, text):
        # Handles **bold** and *italic* text
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for token in tokens:
            if token.startswith('**') and token.endswith('**'):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*'):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                paragraph.add_run(token)

if __name__ == "__main__":
    app = ClaudeApp()
    app.mainloop()